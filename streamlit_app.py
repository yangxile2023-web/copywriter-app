# -*- coding: utf-8 -*-
import streamlit as st
import re
import io
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 页面配置
st.set_page_config(page_title="文案助手", layout="wide")

# 简单CSS
st.markdown("""
<style>
    .main { max-width: 1000px; margin: 0 auto; padding: 20px; }
    .stButton>button[kind="primary"] { 
        background: linear-gradient(90deg, #667eea, #764ba2) !important;
        border: none !important; border-radius: 8px !important;
        padding: 12px 24px !important; font-weight: 600 !important;
    }
    .result-box {
        background: #f8fafc; border-radius: 8px; padding: 16px;
        margin: 8px 0; border-left: 3px solid #667eea;
    }
</style>
""", unsafe_allow_html=True)

# 配置
INDUSTRIES = ["餐饮", "木作定制", "工厂/制造", "彩票店", "酒店/民宿", "通用"]
CONTENT_TYPES = ["干货避坑", "人设故事", "细节特写", "认知反转"]
KIMI_KEY = "sk-IA6qyNJFSYC8UB9RHnGHsgz24VWSrKalSnd5nTTbJNiqQ2uu"

# 确保session state初始化
for key in ['items', 'generating']:
    if key not in st.session_state:
        st.session_state[key] = [] if key == 'items' else False

@st.cache_resource
def get_client():
    return OpenAI(api_key=KIMI_KEY, base_url="https://api.moonshot.cn/v1")

def generate_one(raw_data, idx, length):
    """生成单条文案"""
    try:
        name = re.search(r'出镜称呼[：:]\s*(\S+)', raw_data)
        name = name.group(1) if name else "老板"
        min_w, max_w = (150, 180) if length == "short" else (200, 300)
        ctype = CONTENT_TYPES[(idx - 1) % 4]
        
        client = get_client()
        resp = client.chat.completions.create(
            model="moonshot-v1-8k",
            messages=[{
                "role": "user",
                "content": f"创作第{idx}条{ctype}文案。称呼：{name} 资料：{raw_data[:400]} 要求:字数{min_w}-{max_w} 禁止自我介绍"
            }],
            max_tokens=400, temperature=0.8, timeout=25
        )
        content = resp.choices[0].message.content.strip().strip('"')
        wc = len(content.replace(' ', '').replace('\n', ''))
        return {
            'idx': idx, 'content': content, 'wc': wc,
            'ok': min_w <= wc <= max_w, 'type': ctype
        }
    except Exception as e:
        return {'idx': idx, 'content': f"错误:{str(e)[:20]}", 'wc': 0, 'ok': False, 'type': CONTENT_TYPES[(idx-1)%4]}

def make_word(items, industry, raw_data=""):
    """生成专业格式的Word文档"""
    doc = Document()
    
    # 提取客户称呼作为文件名的一部分
    name_match = re.search(r'出镜称呼[：:]\s*(\S+)', raw_data)
    customer_name = name_match.group(1) if name_match else "客户"
    shop_match = re.search(r'店铺[名称]*[：:]\s*(\S+)', raw_data)
    shop_name = shop_match.group(1) if shop_match else ""
    
    # ===== 封面页 =====
    title = doc.add_heading('短视频文案方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 客户信息表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    info_data = [
        ('客户名称', customer_name),
        ('店铺/品牌', shop_name or industry),
        ('文案数量', f'{len(items)} 条'),
        ('生成日期', datetime.now().strftime('%Y年%m月%d日'))
    ]
    
    for i, (key, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        # 设置单元格样式
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ===== 统计概览 =====
    doc.add_heading('一、文案统计', level=1)
    
    # 按类型统计
    type_counts = {}
    for item in items:
        t = item['type']
        type_counts[t] = type_counts.get(t, 0) + 1
    
    stats_table = doc.add_table(rows=len(type_counts)+1, cols=3)
    stats_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['文案类型', '数量', '占比']
    for i, header in enumerate(headers):
        stats_table.rows[0].cells[i].text = header
    
    # 数据
    for idx, (ctype, count) in enumerate(type_counts.items(), 1):
        row = stats_table.rows[idx]
        row.cells[0].text = ctype
        row.cells[1].text = str(count)
        row.cells[2].text = f'{count/len(items)*100:.1f}%'
    
    doc.add_page_break()
    
    # ===== 分类展示 =====
    doc.add_heading('二、文案详情', level=1)
    
    # 按类型分组展示
    for ctype in CONTENT_TYPES:
        type_items = [i for i in items if i['type'] == ctype]
        if not type_items:
            continue
        
        # 类型标题
        doc.add_heading(f'【{ctype}】', level=2)
        
        for item in type_items:
            # 编号和状态
            p = doc.add_paragraph()
            run = p.add_run(f"文案 #{item['idx']} ")
            run.bold = True
            run.font.size = Pt(11)
            
            # 质量标识
            status = "【优质】" if item['ok'] else "【需优化】"
            status_run = p.add_run(status)
            status_run.font.size = Pt(9)
            status_run.font.color.rgb = RGBColor(34, 197, 94) if item['ok'] else RGBColor(245, 158, 11)
            
            # 文案内容
            content_p = doc.add_paragraph(item['content'])
            content_p.paragraph_format.line_spacing = 1.5
            content_p.paragraph_format.space_after = Pt(6)
            
            # 字数信息
            info_p = doc.add_paragraph(f"字数统计: {item['wc']} 字")
            info_p.runs[0].font.size = Pt(9)
            info_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # 间距
    
    # ===== 使用说明 =====
    doc.add_page_break()
    doc.add_heading('三、使用说明', level=1)
    
    instructions = [
        "1. 本方案共包含30条短视频文案，分为4种类型：",
        "   • 干货避坑：揭露行业内幕，提供避坑指南",
        "   • 人设故事：讲述老板个人经历，打造人设",
        "   • 细节特写：展示产品/工艺的具体细节",
        "   • 认知反转：颠覆常识，打破认知",
        "",
        "2. 文案可直接用于短视频拍摄，建议：",
        "   • 根据实际场景选择合适的文案",
        "   • 可适当调整文案中的称呼和细节",
        "   • 优质文案优先使用，需优化文案可调整后再用",
        "",
        "3. 拍摄建议：",
        "   • 开头3秒必须有吸引力（利益/冲突/悬念）",
        "   • 避免自我介绍和店铺地址",
        "   • 多用具体数字和细节增加真实感",
    ]
    
    for line in instructions:
        doc.add_paragraph(line)
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    footer = doc.add_paragraph('晓牧传媒 · 内容创作系统')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ========== UI ==========
st.title("📝 晓牧传媒文案助手")
st.caption("AI生成30条短视频文案")

# 配置
st.subheader("配置")
c1, c2 = st.columns(2)
with c1:
    industry = st.selectbox("行业", INDUSTRIES, index=5)
with c2:
    length_opt = st.radio("长度", ["短文案(150-180)", "长文案(200-300)"], horizontal=True)

# 输入
st.subheader("客户资料")
raw = st.text_area("输入", height=120, placeholder="请输入客户资料：出镜称呼、店铺名称、主营业务、真实故事...")
count = len(raw.replace(' ', '').replace('\n', ''))
st.caption(f"已输入 {count} 字")

# 生成按钮
if st.button("🚀 生成30条文案", type="primary", disabled=st.session_state.generating):
    if count < 30:
        st.error("请至少输入30字")
    else:
        st.session_state.generating = True
        length = "short" if "短" in length_opt else "long"
        
        progress = st.progress(0)
        results = []
        
        for i in range(30):
            result = generate_one(raw, i + 1, length)
            results.append(result)
            progress.progress((i + 1) / 30)
        
        st.session_state.items = results
        st.session_state.raw_data = raw
        st.session_state.generating = False
        st.success(f"✅ 生成完成！{sum(1 for r in results if r['ok'])} 条优质文案")
        st.rerun()

# 显示结果
items = st.session_state.get('items', [])
if isinstance(items, list) and len(items) > 0:
    st.divider()
    st.subheader(f"生成结果 ({len(items)}条)")
    
    # 统计
    ok_num = sum(1 for i in items if i.get('ok'))
    cols = st.columns(4)
    cols[0].metric("总数", len(items))
    cols[1].metric("优质", ok_num)
    cols[2].metric("需优化", len(items) - ok_num)
    
    # 下载按钮
    with cols[3]:
        try:
            raw_data = st.session_state.get('raw_data', '')
            word_data = make_word(items, industry, raw_data)
            # 提取客户名用于文件名
            name_match = re.search(r'出镜称呼[：:]\s*(\S+)', raw_data)
            customer_name = name_match.group(1) if name_match else industry
            st.download_button(
                "📥 下载Word",
                word_data,
                file_name=f"文案_{customer_name}_{len(items)}条.docx",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Word生成失败: {e}")
    
    # 清空按钮
    if st.button("🗑️ 清空结果"):
        st.session_state.items = []
        st.rerun()
    
    # 文案列表
    st.subheader("文案详情")
    for item in items:
        status = "✅" if item.get('ok') else "⚠️"
        with st.expander(f"{status} 文案 #{item['idx']} [{item['type']}] - {item['wc']}字"):
            st.write(item['content'])
            st.button("复制", key=f"copy_{item['idx']}")
